"""
Document Ingestion Pipeline
Processes various document formats for the knowledge base
"""

import asyncio
import logging
import hashlib
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
import json
import re

# Document processing libraries
import PyPDF2
import docx
import markdown
import csv
import openpyxl
from bs4 import BeautifulSoup

from agi.core.database import get_db_manager
from agi.knowledge.hybrid_search import get_hybrid_search
from agi.config.agi_config import get_config

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types"""
    PDF = "pdf"
    WORD = "docx"
    TEXT = "txt"
    MARKDOWN = "md"
    HTML = "html"
    CSV = "csv"
    EXCEL = "xlsx"
    JSON = "json"
    CODE = "code"
    UNKNOWN = "unknown"


class ProcessingStatus(Enum):
    """Document processing status"""
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    INDEXED = "indexed"


@dataclass
class DocumentChunk:
    """Represents a chunk of a document"""
    id: str
    document_id: str
    content: str
    chunk_index: int
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None


@dataclass
class ProcessedDocument:
    """Represents a processed document"""
    id: str
    source_path: str
    document_type: DocumentType
    title: str
    content: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    processing_status: ProcessingStatus
    created_at: datetime = field(default_factory=datetime.utcnow)
    indexed_at: Optional[datetime] = None
    error_message: Optional[str] = None


class DocumentProcessor:
    """Base class for document processors"""

    async def process(self, file_path: str, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process document and extract text and metadata"""
        raise NotImplementedError


class PDFProcessor(DocumentProcessor):
    """Process PDF documents"""

    async def process(self, file_path: str, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from PDF"""
        try:
            import io
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_content = []
            metadata = {
                "num_pages": len(pdf_reader.pages),
                "info": pdf_reader.metadata if pdf_reader.metadata else {}
            }

            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                text_content.append(page_text)

            return "\n\n".join(text_content), metadata

        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise


class WordProcessor(DocumentProcessor):
    """Process Word documents"""

    async def process(self, file_path: str, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from Word document"""
        try:
            import io
            doc = docx.Document(io.BytesIO(content))

            text_content = []
            metadata = {
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables)
            }

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                text_content.append("\n".join(table_text))

            return "\n\n".join(text_content), metadata

        except Exception as e:
            logger.error(f"Word processing failed: {e}")
            raise


class TextProcessor(DocumentProcessor):
    """Process plain text documents"""

    async def process(self, file_path: str, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file"""
        try:
            text = content.decode('utf-8', errors='ignore')
            metadata = {
                "num_lines": len(text.split('\n')),
                "num_words": len(text.split())
            }
            return text, metadata

        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            raise


class MarkdownProcessor(DocumentProcessor):
    """Process Markdown documents"""

    async def process(self, file_path: str, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from Markdown"""
        try:
            md_text = content.decode('utf-8', errors='ignore')

            # Extract metadata from frontmatter if present
            metadata = {}
            if md_text.startswith('---'):
                parts = md_text.split('---', 2)
                if len(parts) >= 3:
                    # Parse YAML frontmatter
                    frontmatter = parts[1].strip()
                    for line in frontmatter.split('\n'):
                        if ':' in line:
                            key, value = line.split(':', 1)
                            metadata[key.strip()] = value.strip()
                    md_text = parts[2]

            # Convert to plain text
            html = markdown.markdown(md_text)
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()

            metadata['has_code_blocks'] = '```' in md_text
            metadata['num_headers'] = len(re.findall(r'^#+\s', md_text, re.MULTILINE))

            return text, metadata

        except Exception as e:
            logger.error(f"Markdown processing failed: {e}")
            raise


class HTMLProcessor(DocumentProcessor):
    """Process HTML documents"""

    async def process(self, file_path: str, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from HTML"""
        try:
            html = content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html, 'html.parser')

            # Remove script and style elements
            for element in soup(['script', 'style']):
                element.decompose()

            # Extract metadata
            metadata = {}
            if soup.title:
                metadata['title'] = soup.title.string

            meta_tags = soup.find_all('meta')
            for tag in meta_tags:
                if tag.get('name'):
                    metadata[tag['name']] = tag.get('content', '')

            # Extract text
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)

            return text, metadata

        except Exception as e:
            logger.error(f"HTML processing failed: {e}")
            raise


class CSVProcessor(DocumentProcessor):
    """Process CSV files"""

    async def process(self, file_path: str, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract data from CSV"""
        try:
            import io
            text = content.decode('utf-8', errors='ignore')
            csv_file = io.StringIO(text)
            reader = csv.DictReader(csv_file)

            rows = list(reader)
            metadata = {
                "num_rows": len(rows),
                "columns": reader.fieldnames if reader.fieldnames else []
            }

            # Convert to text format
            text_content = []
            for row in rows[:1000]:  # Limit to first 1000 rows
                row_text = ", ".join(f"{k}: {v}" for k, v in row.items())
                text_content.append(row_text)

            return "\n".join(text_content), metadata

        except Exception as e:
            logger.error(f"CSV processing failed: {e}")
            raise


class ExcelProcessor(DocumentProcessor):
    """Process Excel files"""

    async def process(self, file_path: str, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract data from Excel"""
        try:
            import io
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True)

            metadata = {
                "num_sheets": len(wb.sheetnames),
                "sheet_names": wb.sheetnames
            }

            text_content = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_content.append(f"Sheet: {sheet_name}")

                for row in sheet.iter_rows(max_row=1000, values_only=True):
                    row_text = " | ".join(str(cell) for cell in row if cell)
                    if row_text:
                        text_content.append(row_text)

            return "\n\n".join(text_content), metadata

        except Exception as e:
            logger.error(f"Excel processing failed: {e}")
            raise


class CodeProcessor(DocumentProcessor):
    """Process source code files"""

    async def process(self, file_path: str, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract and analyze source code"""
        try:
            code = content.decode('utf-8', errors='ignore')

            # Detect language from extension
            ext = Path(file_path).suffix.lower()
            language_map = {
                '.py': 'python',
                '.js': 'javascript',
                '.ts': 'typescript',
                '.java': 'java',
                '.cpp': 'cpp',
                '.c': 'c',
                '.go': 'go',
                '.rs': 'rust',
                '.rb': 'ruby',
                '.php': 'php'
            }

            language = language_map.get(ext, 'unknown')

            # Extract metadata
            metadata = {
                "language": language,
                "num_lines": len(code.split('\n')),
                "has_comments": '//' in code or '#' in code or '/*' in code
            }

            # Add language context to text
            text = f"Programming Language: {language}\n\n{code}"

            return text, metadata

        except Exception as e:
            logger.error(f"Code processing failed: {e}")
            raise


class DocumentIngestionPipeline:
    """
    Complete document ingestion pipeline
    Features: format detection, chunking, indexing, deduplication
    """

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.processors = {
            DocumentType.PDF: PDFProcessor(),
            DocumentType.WORD: WordProcessor(),
            DocumentType.TEXT: TextProcessor(),
            DocumentType.MARKDOWN: MarkdownProcessor(),
            DocumentType.HTML: HTMLProcessor(),
            DocumentType.CSV: CSVProcessor(),
            DocumentType.EXCEL: ExcelProcessor(),
            DocumentType.CODE: CodeProcessor()
        }
        self.processed_documents = {}

    async def ingest_document(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """
        Ingest a document into the knowledge base

        Args:
            file_path: Path to the document
            metadata: Optional metadata to attach

        Returns:
            ProcessedDocument object
        """
        try:
            # Generate document ID
            doc_id = self._generate_document_id(file_path)

            # Check if already processed
            if await self._is_duplicate(doc_id):
                logger.info(f"Document {file_path} already processed")
                return self.processed_documents.get(doc_id)

            # Detect document type
            doc_type = self._detect_document_type(file_path)

            # Create processed document
            processed_doc = ProcessedDocument(
                id=doc_id,
                source_path=file_path,
                document_type=doc_type,
                title=Path(file_path).name,
                content="",
                chunks=[],
                metadata=metadata or {},
                processing_status=ProcessingStatus.PROCESSING
            )

            # Read file content
            with open(file_path, 'rb') as f:
                content = f.read()

            # Process document based on type
            if doc_type in self.processors:
                processor = self.processors[doc_type]
                text_content, doc_metadata = await processor.process(file_path, content)

                processed_doc.content = text_content
                processed_doc.metadata.update(doc_metadata)

                # Create chunks
                chunks = await self._create_chunks(processed_doc)
                processed_doc.chunks = chunks

                # Index document
                await self._index_document(processed_doc)

                processed_doc.processing_status = ProcessingStatus.INDEXED
                processed_doc.indexed_at = datetime.utcnow()

            else:
                processed_doc.processing_status = ProcessingStatus.FAILED
                processed_doc.error_message = f"Unsupported document type: {doc_type}"

            # Store in database
            await self._store_document(processed_doc)

            # Cache processed document
            self.processed_documents[doc_id] = processed_doc

            logger.info(f"Successfully ingested document: {file_path}")
            return processed_doc

        except Exception as e:
            logger.error(f"Document ingestion failed for {file_path}: {e}")
            raise

    async def ingest_batch(
        self,
        file_paths: List[str],
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[ProcessedDocument]:
        """
        Ingest multiple documents in batch

        Args:
            file_paths: List of document paths
            metadata: Optional metadata for all documents

        Returns:
            List of ProcessedDocument objects
        """
        tasks = []
        for file_path in file_paths:
            task = self.ingest_document(file_path, metadata)
            tasks.append(task)

        results = await asyncio.gather(*tasks, return_exceptions=True)

        processed = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"Failed to process {file_paths[i]}: {result}")
            else:
                processed.append(result)

        return processed

    async def update_document(
        self,
        file_path: str,
        force: bool = False
    ) -> ProcessedDocument:
        """Update an existing document"""
        doc_id = self._generate_document_id(file_path)

        # Remove old version if exists
        if doc_id in self.processed_documents or force:
            await self._remove_document(doc_id)

        # Re-ingest
        return await self.ingest_document(file_path)

    def _detect_document_type(self, file_path: str) -> DocumentType:
        """Detect document type from file extension"""
        ext = Path(file_path).suffix.lower()

        type_map = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.WORD,
            '.doc': DocumentType.WORD,
            '.txt': DocumentType.TEXT,
            '.md': DocumentType.MARKDOWN,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
            '.csv': DocumentType.CSV,
            '.xlsx': DocumentType.EXCEL,
            '.xls': DocumentType.EXCEL,
            '.json': DocumentType.JSON,
            '.py': DocumentType.CODE,
            '.js': DocumentType.CODE,
            '.ts': DocumentType.CODE,
            '.java': DocumentType.CODE,
            '.cpp': DocumentType.CODE,
            '.c': DocumentType.CODE,
            '.go': DocumentType.CODE,
            '.rs': DocumentType.CODE
        }

        return type_map.get(ext, DocumentType.UNKNOWN)

    def _generate_document_id(self, file_path: str) -> str:
        """Generate unique document ID"""
        path_str = str(Path(file_path).absolute())
        return hashlib.md5(path_str.encode()).hexdigest()

    async def _is_duplicate(self, doc_id: str) -> bool:
        """Check if document is already processed"""
        if doc_id in self.processed_documents:
            return True

        # Check database
        db = await get_db_manager()
        result = await db.fetchone(
            "SELECT id FROM agi.document_chunks WHERE document_id = $1 LIMIT 1",
            doc_id
        )

        return result is not None

    async def _create_chunks(self, document: ProcessedDocument) -> List[DocumentChunk]:
        """Create chunks from document content"""
        chunks = []
        content = document.content

        # Simple chunking by character count with overlap
        words = content.split()
        current_chunk = []
        current_size = 0

        for i, word in enumerate(words):
            current_chunk.append(word)
            current_size += len(word) + 1

            if current_size >= self.chunk_size:
                # Create chunk
                chunk_text = " ".join(current_chunk)
                chunk_id = f"{document.id}_chunk_{len(chunks)}"

                chunk = DocumentChunk(
                    id=chunk_id,
                    document_id=document.id,
                    content=chunk_text,
                    chunk_index=len(chunks),
                    metadata={
                        "source": document.source_path,
                        "type": document.document_type.value,
                        "chunk_size": len(chunk_text)
                    }
                )
                chunks.append(chunk)

                # Overlap for next chunk
                overlap_words = int(self.chunk_overlap * len(current_chunk) / current_size)
                current_chunk = current_chunk[-overlap_words:] if overlap_words > 0 else []
                current_size = sum(len(w) + 1 for w in current_chunk)

        # Add remaining content as last chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunk_id = f"{document.id}_chunk_{len(chunks)}"

            chunk = DocumentChunk(
                id=chunk_id,
                document_id=document.id,
                content=chunk_text,
                chunk_index=len(chunks),
                metadata={
                    "source": document.source_path,
                    "type": document.document_type.value,
                    "chunk_size": len(chunk_text)
                }
            )
            chunks.append(chunk)

        logger.info(f"Created {len(chunks)} chunks for document {document.id}")
        return chunks

    async def _index_document(self, document: ProcessedDocument):
        """Index document in hybrid search"""
        try:
            hybrid_search = await get_hybrid_search()

            # Prepare documents for indexing
            search_docs = []
            for chunk in document.chunks:
                search_doc = {
                    'content': chunk.content,
                    'metadata': {
                        **chunk.metadata,
                        **document.metadata,
                        'title': document.title,
                        'indexed_at': datetime.utcnow().isoformat()
                    },
                    'document_id': document.id,
                    'chunk_id': chunk.id
                }
                search_docs.append(search_doc)

            # Index in hybrid search
            await hybrid_search.index_documents(search_docs)

            logger.info(f"Indexed {len(search_docs)} chunks for document {document.id}")

        except Exception as e:
            logger.error(f"Failed to index document {document.id}: {e}")
            raise

    async def _store_document(self, document: ProcessedDocument):
        """Store document in database"""
        try:
            db = await get_db_manager()

            # Store chunks
            for chunk in document.chunks:
                await db.execute(
                    """
                    INSERT INTO agi.document_chunks
                    (id, document_id, content, chunk_index, metadata, created_at)
                    VALUES ($1, $2, $3, $4, $5, $6)
                    ON CONFLICT (id) DO UPDATE SET
                        content = $3,
                        chunk_index = $4,
                        metadata = $5
                    """,
                    chunk.id,
                    chunk.document_id,
                    chunk.content,
                    chunk.chunk_index,
                    json.dumps(chunk.metadata),
                    datetime.utcnow()
                )

            logger.info(f"Stored {len(document.chunks)} chunks in database")

        except Exception as e:
            logger.error(f"Failed to store document {document.id}: {e}")
            raise

    async def _remove_document(self, doc_id: str):
        """Remove document from system"""
        try:
            # Remove from database
            db = await get_db_manager()
            await db.execute(
                "DELETE FROM agi.document_chunks WHERE document_id = $1",
                doc_id
            )
            await db.execute(
                "DELETE FROM agi.document_embeddings WHERE document_id = $1",
                doc_id
            )

            # Remove from cache
            if doc_id in self.processed_documents:
                del self.processed_documents[doc_id]

            # Remove from search index
            hybrid_search = await get_hybrid_search()
            await hybrid_search.delete_document(doc_id)

            logger.info(f"Removed document {doc_id}")

        except Exception as e:
            logger.error(f"Failed to remove document {doc_id}: {e}")

    def get_stats(self) -> Dict[str, Any]:
        """Get ingestion pipeline statistics"""
        return {
            "processed_documents": len(self.processed_documents),
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "supported_types": [t.value for t in DocumentType]
        }


# Global ingestion pipeline
_ingestion_pipeline: Optional[DocumentIngestionPipeline] = None


async def get_ingestion_pipeline() -> DocumentIngestionPipeline:
    """Get singleton ingestion pipeline instance"""
    global _ingestion_pipeline
    if _ingestion_pipeline is None:
        config = get_config()
        _ingestion_pipeline = DocumentIngestionPipeline(
            chunk_size=config.rag.chunk_size,
            chunk_overlap=config.rag.chunk_overlap
        )
    return _ingestion_pipeline